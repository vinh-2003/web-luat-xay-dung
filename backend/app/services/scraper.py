import asyncio
import os
import aiofiles
import pdfplumber
import docx
from playwright.async_api import async_playwright
from datetime import datetime
from sqlalchemy.orm import Session
import traceback
from app.db.database import SessionLocal
from app.models.document import Document, DocumentChunk
from app.models.scraper import ScraperJob
from app.services.gemini import GeminiService

class LuatXayDungScraper:
    def __init__(self):
        self.gemini_service = GeminiService()

    async def _check_job_status(self, db: Session, job_id: int):
        if not job_id:
            return True
        job = db.query(ScraperJob).filter(ScraperJob.id == job_id).first()
        if not job:
            return False
        
        if job.status == "STOPPED":
            print("Job STOPPED. Exiting scraper loop...")
            return False
            
        # Kiểm tra PAUSED và timeout
        pause_start = datetime.now()
        while job.status == "PAUSED":
            elapsed = (datetime.now() - pause_start).total_seconds() / 60
            if elapsed > (job.timeout_minutes or 60):
                print("Job PAUSED quá thời gian (timeout). Tự động STOP.")
                job.status = "STOPPED"
                db.commit()
                return False
                
            print("Job PAUSED. Waiting for RESUME...")
            await asyncio.sleep(5)
            db.refresh(job)
            if job.status in ["STOPPED", "FAILED", "COMPLETED"]:
                return False
                
        return True

    def _chunk_text(self, text: str, chunk_size: int = 2500, overlap_words: int = 50) -> list[str]:
        words = text.split()
        chunks = []
        i = 0
        while i < len(words):
            current_chunk = []
            current_length = 0
            while i < len(words) and current_length < chunk_size:
                current_length += len(words[i]) + 1
                current_chunk.append(words[i])
                i += 1
            chunks.append(" ".join(current_chunk))
            
            # Nếu chưa đến cuối, lùi lại overlap_words để tạo vùng giao thoa (overlap)
            if i < len(words):
                i = max(0, i - overlap_words)
        return chunks

    def _extract_text_from_file(self, file_path: str) -> str:
        text = ""
        try:
            if file_path.lower().endswith('.pdf'):
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
            elif file_path.lower().endswith('.docx'):
                doc = docx.Document(file_path)
                for para in doc.paragraphs:
                    text += para.text + "\n"
        except Exception as e:
            print(f"Lỗi extract file {file_path}: {e}")
        return text

    async def _download_and_extract_file(self, page, locator) -> str:
        text = ""
        try:
            href = await locator.get_attribute("href")
            if href and href != "#" and not href.startswith("javascript"):
                # Handle relative links
                if not href.startswith("http"):
                    if href.startswith("/"):
                        href = "https://moc.gov.vn" + href
                    else:
                        href = "https://moc.gov.vn/vn/pages/" + href

                print(f"Đang tải nội dung tệp: {href}")
                response = await page.request.get(href)
                if response.ok:
                    body = await response.body()
                    ext = href.split(".")[-1].lower()
                    if "?" in ext:
                        ext = ext.split("?")[0]
                    if not ext in ['pdf', 'doc', 'docx']:
                        ext = "pdf"
                        
                    import uuid
                    file_name = f"temp_{uuid.uuid4().hex}.{ext}"
                    dest_path = os.path.join(os.getcwd(), file_name)
                    
                    with open(dest_path, "wb") as f:
                        f.write(body)
                        
                    text = self._extract_text_from_file(dest_path)
                    if os.path.exists(dest_path):
                        os.remove(dest_path)
                    return text
                else:
                    print(f"Lỗi tải HTTP status {response.status}")

            # Fallback cách tải tự động của trình duyệt
            async with page.expect_download(timeout=10000) as download_info:
                await locator.click()
            download = await download_info.value
            file_name = download.suggested_filename
            if not (file_name.lower().endswith('.pdf') or file_name.lower().endswith('.docx') or file_name.lower().endswith('.doc')):
                print(f"Bỏ qua định dạng tệp: {file_name}")
                await download.cancel()
                return ""
            
            dest_path = os.path.join(os.getcwd(), file_name)
            await download.save_as(dest_path)
            
            # Extract text
            text = self._extract_text_from_file(dest_path)
            
            # Clean up
            if os.path.exists(dest_path):
                os.remove(dest_path)
        except Exception as e:
            print(f"Lỗi khi tải hoặc đọc tệp: {e}")
        return text

    def _save_document(self, db: Session, doc_data: dict, job_id: int = None) -> str:
        """Trả về 'added' hoặc 'updated' hoặc 'skipped'"""
        # Tiền xử lý ngày tháng để tránh lỗi PostgreSQL
        for date_field in ["issue_date", "effective_date"]:
            val = doc_data.get(date_field)
            if not val or val == "N/A" or val.lower() == "n/a" or "ngày" in val.lower():
                doc_data[date_field] = None
            else:
                try:
                    datetime.strptime(val, "%Y-%m-%d")
                except:
                    doc_data[date_field] = None
                    
        # Find by document_number if possible, otherwise by title
        query_filter = Document.document_number == doc_data.get("document_number") if doc_data.get("document_number") != "N/A" else Document.title == doc_data["title"]
        existing_doc = db.query(Document).filter(query_filter).first()
        
        if existing_doc:
            print(f"Cập nhật trạng thái cho: {doc_data['title']}...")
            updated = False
            for key in ["status", "issue_date", "effective_date", "document_type", "signer", "issuing_body"]:
                new_val = doc_data.get(key)
                if new_val is not None and new_val != "N/A" and getattr(existing_doc, key) != new_val:
                    setattr(existing_doc, key, new_val)
                    updated = True
            
            new_content = doc_data.get("content")
            if new_content and len(new_content) > 100:
                old_content = existing_doc.content or ""
                if len(old_content) < 100:
                    print("Cập nhật content từ tệp đính kèm...")
                    existing_doc.content = new_content
                    updated = True
                    # Xoá chunks cũ
                    db.query(DocumentChunk).filter(DocumentChunk.document_id == existing_doc.id).delete()
                    # Thêm chunks mới
                    chunks = self._chunk_text(new_content)
                    
                    chunk_texts = []
                    for original_chunk_text in chunks:
                        chunk_text = f"[Văn bản: {existing_doc.title}]\n[Số hiệu: {existing_doc.document_number}]\n\n{original_chunk_text}"
                        chunk_texts.append(chunk_text)
                        
                    batch_size = 10
                    for i in range(0, len(chunk_texts), batch_size):
                        batch_texts = chunk_texts[i:i + batch_size]
                        try:
                            batch_embeddings = self.gemini_service.generate_embeddings_batch(batch_texts)
                        except Exception as e:
                            print(f"Lỗi khi sinh embedding batch: {e}")
                            batch_embeddings = [None] * len(batch_texts)
                            
                        for j, embed_vector in enumerate(batch_embeddings):
                            chunk_obj = DocumentChunk(document_id=existing_doc.id, chunk_text=batch_texts[j], embedding=embed_vector)
                            db.add(chunk_obj)

            if updated:
                db.commit()
                print(f"Đã cập nhật thông tin thành công cho: {doc_data['title']}.")
                return "updated"
            else:
                print(f"Không có thay đổi thông tin cho: {doc_data['title']}.")
            return "skipped"

        print(f"Đang lưu mới: {doc_data['title']}...")
        new_doc = Document(
            document_number=doc_data["document_number"],
            title=doc_data["title"],
            issue_date=doc_data.get("issue_date"),
            effective_date=doc_data.get("effective_date"),
            status=doc_data.get("status", "N/A"),
            document_type=doc_data.get("document_type"),
            signer=doc_data.get("signer"),
            issuing_body=doc_data.get("issuing_body"),
            content=doc_data["content"],
            origin_url=doc_data["origin_url"]
        )
        db.add(new_doc)
        db.commit()
        db.refresh(new_doc)

        chunks = self._chunk_text(new_doc.content)
        chunk_texts = []
        for original_chunk_text in chunks:
            chunk_text = f"[Văn bản: {new_doc.title}]\n[Số hiệu: {new_doc.document_number}]\n\n{original_chunk_text}"
            chunk_texts.append(chunk_text)
            
        batch_size = 10
        for i in range(0, len(chunk_texts), batch_size):
            batch_texts = chunk_texts[i:i + batch_size]
            try:
                batch_embeddings = self.gemini_service.generate_embeddings_batch(batch_texts)
            except Exception as e:
                print(f"Lỗi khi sinh embedding batch: {e}")
                batch_embeddings = [None] * len(batch_texts)
                
            for j, embed_vector in enumerate(batch_embeddings):
                chunk_obj = DocumentChunk(
                    document_id=new_doc.id,
                    chunk_text=batch_texts[j],
                    embedding=embed_vector
                )
                db.add(chunk_obj)

        db.commit()
        print(f"Đã lưu thành công: {doc_data['title']} với {len(chunks)} chunks.")
        return "added"

    def _parse_properties(self, props_text: str, doc_data: dict):
        lines = [line.strip() for line in props_text.split('\n') if line.strip()]
        for i, line in enumerate(lines):
            if line == "Số hiệu" and i + 1 < len(lines): doc_data["document_number"] = lines[i+1]
            elif line == "Loại văn bản" and i + 1 < len(lines): doc_data["document_type"] = lines[i+1]
            elif line == "Ngày ban hành" and i + 1 < len(lines): doc_data["issue_date"] = lines[i+1]
            elif line == "Ngày có hiệu lực" and i + 1 < len(lines): doc_data["effective_date"] = lines[i+1]
            elif line == "Tình trạng hiệu lực" and i + 1 < len(lines): doc_data["status"] = lines[i+1]
            elif line == "Cơ quan ban hành" and i + 1 < len(lines): doc_data["issuing_body"] = lines[i+1]
            elif line == "Người ký" and i + 1 < len(lines): doc_data["signer"] = lines[i+1]
        
        if doc_data.get("document_number") == "N/A" and doc_data.get("title") != "N/A":
            import re
            m = re.search(r'([0-9]+/[0-9A-Za-z]+/[0-9A-Za-z.-]+)', doc_data["title"])
            if m:
                doc_data["document_number"] = m.group(1)
                
        # Fix date formats
        from datetime import datetime
        for date_field in ["issue_date", "effective_date"]:
            if doc_data.get(date_field) and doc_data[date_field] != "N/A":
                try:
                    dt = datetime.strptime(doc_data[date_field].strip(), "%d/%m/%Y")
                    doc_data[date_field] = dt.strftime("%Y-%m-%d")
                except ValueError:
                    pass

    async def scrape_vbpl(self, job_id: int = None, mode: str = "full"):
        """
        Cào dữ liệu từ vbpl.vn với cơ chế Hybrid Search
        """
        db = SessionLocal()
        async with async_playwright() as p:
            browser = await p.chromium.launch(
                headless=True, 
                args=['--start-maximized', '--disable-blink-features=AutomationControlled'],
                proxy={
                    "server": "http://103.129.127.244:8088"
                }
            )
            context = await browser.new_context(
                no_viewport=True,
                user_agent="Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
                ignore_https_errors=True
            )
            
            try:
                keywords = ["xây dựng", "kiến trúc", "quy hoạch", "nhà ở", "bất động sản", "đấu thầu", "vật liệu xây dựng", "hạ tầng", "công trình"]

                for keyword in keywords:
                    if not await self._check_job_status(db, job_id):
                        break

                    page = await context.new_page()
                    
                    print(f"=== TÌM KIẾM TỪ KHÓA: {keyword} ===")
                    await page.goto("https://vbpl.vn/", timeout=60000, wait_until="domcontentloaded")
                    await page.wait_for_load_state("networkidle")
                    await page.wait_for_timeout(2000)
                    
                    # Điền từ khóa
                    inputs = await page.locator("input[type='text']").all()
                    for inp in inputs:
                        ph = await inp.get_attribute("placeholder") or ""
                        if "Tìm kiếm" in ph or "nhập" in ph.lower() or "văn bản" in ph.lower():
                            await inp.fill(keyword)
                            break
                            
                    # Chọn "Nội dung" radio button
                    content_radio = page.locator("input[type='radio'][value='2']") # vbpl.vn thường dùng value='2' cho Nội dung, hoặc dùng text
                    if await content_radio.count() > 0:
                        await content_radio.check(force=True)
                    else:
                        # Fallback bằng click text
                        content_label = page.locator("label:has-text('Nội dung')").first
                        if await content_label.count() > 0:
                            await content_label.click()
                    
                    # Click Tìm kiếm
                    search_btn = page.locator("button:has-text('Tìm kiếm')").first
                    if await search_btn.count() > 0:
                        await search_btn.click()
                    else:
                        await page.keyboard.press("Enter")
                        
                    await page.wait_for_load_state("networkidle")
                    await page.wait_for_timeout(5000)
                    
                    # Sắp xếp "Từ mới đến cũ"
                    try:
                        print("Thực hiện sắp xếp 'Từ mới đến cũ'...")
                        
                        # 1. Mở menu dropdown (Giá trị mặc định thường là 'Từ khóa tìm kiếm')
                        try:
                            btn = page.locator("text='Từ khóa tìm kiếm'").first
                            if await btn.count() > 0:
                                await btn.click(force=True)
                            else:
                                await page.locator("text='Sắp xếp theo'").locator("xpath=..").click(force=True)
                            await page.wait_for_timeout(1000)
                        except:
                            pass
                            
                        # 2. Click vào 'Ngày ban hành'
                        try:
                            # Lấy thẻ chứa text chính xác để tránh nhầm label
                            ngay_ban_hanh = page.locator("text='Ngày ban hành'").last
                            if await ngay_ban_hanh.count() > 0:
                                await ngay_ban_hanh.click(force=True)
                                await page.wait_for_timeout(1000)
                        except:
                            pass
                            
                        # 3. Click 'Từ mới đến cũ'
                        try:
                            new_to_old = page.locator("text='Từ mới đến cũ'").last
                            if await new_to_old.count() > 0:
                                await new_to_old.click(force=True)
                                await page.wait_for_load_state("networkidle")
                                await page.wait_for_timeout(4000)
                        except:
                            pass
                            
                        # 4. Fallback bằng JS: Chỉ click chính xác thẻ a hoặc li có nội dung chuẩn (tránh click nhầm div to)
                        await page.evaluate("""() => {
                            let elements = Array.from(document.querySelectorAll('a, li, span'));
                            let target = elements.find(el => el.textContent.trim() === 'Từ mới đến cũ');
                            if (target) {
                                target.click();
                            }
                        }""")
                        
                        await page.wait_for_load_state("networkidle")
                        await page.wait_for_timeout(3000)
                        print("Hoàn tất thao tác sắp xếp.")
                            
                    except Exception as e:
                        print(f"Không thể sắp xếp (bỏ qua): {e}")
                        print(f"Không thể sắp xếp (bỏ qua): {e}")
                    
                    valid_statuses = ["chưa có hiệu lực", "còn hiệu lực", "hết hiệu lực một phần", "ngưng hiệu lực một phần"]
                    page_num = 1
                    
                    while True:
                        if not await self._check_job_status(db, job_id):
                            break

                        if job_id:
                            job = db.query(ScraperJob).filter(ScraperJob.id == job_id).first()
                            if job:
                                job.current_page = page_num
                                db.commit()

                        print(f"--- Đang cào trang {page_num} ---")
                        cards = await page.locator(".DocumentCard_cardContent__5tVRC").all()
                        if not cards:
                            # Try old layout just in case
                            cards = await page.locator("li .item").all()
                            if not cards:
                                print("Không tìm thấy kết quả nào hoặc cấu trúc DOM đã thay đổi.")
                                break
                        
                        print(f"Đã tìm thấy {len(cards)} kết quả trên trang {page_num}.")
                        
                        # Trích xuất tiêu đề
                        page_titles = []
                        for card in cards:
                            title_el = card.locator(".DocumentCard_documentTitle__aE_F_, .title")
                            try:
                                title = await title_el.inner_text()
                                page_titles.append(title)
                            except:
                                page_titles.append("N/A")
                                
                        # Batch check AI
                        try:
                            relevance_results = self.gemini_service.check_documents_relevance_batch(page_titles)
                        except Exception as e:
                            print(f"Lỗi kiểm tra guardrail batch: {e}")
                            relevance_results = [True] * len(page_titles)
                        
                        for i in range(len(cards)):
                            if not await self._check_job_status(db, job_id):
                                break

                            card = cards[i]
                            doc_title = page_titles[i]
                            
                            doc_data = {
                                "title": doc_title,
                                "origin_url": "https://vbpl.vn",
                                "document_number": "N/A",
                                "status": "N/A",
                                "document_type": "N/A",
                                "signer": "N/A",
                                "issuing_body": "N/A",
                                "content": ""
                            }
                            
                            if doc_title == "N/A": continue
                            if not relevance_results[i]:
                                print(f"Bỏ qua: {doc_title} (Không liên quan)")
                                continue
                            
                            title_el = card.locator(".DocumentCard_documentTitle__aE_F_, .title").first

                            print(f"Đang mở chi tiết văn bản: {doc_title}...")
                            
                            try:
                                async with context.expect_page(timeout=15000) as new_page_info:
                                    try:
                                        await title_el.click(timeout=3000, force=True)
                                    except Exception as e:
                                        print(f"Lỗi click (sẽ đợi timeout): {e}")
                                        
                                new_page = await new_page_info.value
                                await new_page.wait_for_load_state("networkidle")
                                await new_page.wait_for_timeout(2000)
                                
                                doc_data["origin_url"] = new_page.url
                                
                                try:
                                    await new_page.wait_for_selector(".ant-tabs-tabpane, #contentBody", timeout=15000)
                                    
                                    content_el = new_page.locator(".ant-tabs-tabpane, #contentBody").first
                                    doc_data["content"] = await content_el.inner_text()
                                    
                                    tab_btn = new_page.locator(".ant-tabs-tab-btn", has_text="Thuộc tính").first
                                    if await tab_btn.count() > 0:
                                        await tab_btn.click()
                                        await new_page.wait_for_timeout(2000)
                                        props_el = new_page.locator(".ant-tabs-tabpane-active").first
                                        props_text = await props_el.inner_text()
                                        self._parse_properties(props_text, doc_data)
                                except Exception as ex:
                                    print(f"Lỗi lấy nội dung chi tiết: {ex}")
                                    
                                await new_page.close()
                                
                            except Exception as e:
                                print(f"Bỏ qua văn bản do lỗi mở tab/timeout: {e}")
                                continue
                                
                            if doc_data["status"].lower() not in valid_statuses and doc_data["status"] != "N/A":
                                print(f"Bỏ qua trạng thái: {doc_data['status']}")
                                continue
                                
                            try:
                                result = self._save_document(db, doc_data, job_id)
                                if job_id:
                                    job = db.query(ScraperJob).filter(ScraperJob.id == job_id).first()
                                    if job:
                                        if result == "added": job.docs_added += 1
                                        elif result == "updated": job.docs_updated += 1
                                        db.commit()
                            except Exception as e:
                                print(f"Lỗi lưu DB (VBPL): {e}")
                                db.rollback()
                                
                        if not await self._check_job_status(db, job_id):
                            break

                        if mode in ["schedule_scrape", "update"] and page_num >= 2:
                            print("Chế độ UPDATE/SCHEDULE: Đã cào xong số trang cho phép (2 trang) cho từ khoá này.")
                            break

                        next_btn = page.locator("li.ant-pagination-next, a.Next").first
                        if await next_btn.count() > 0:
                            is_disabled = await next_btn.get_attribute("aria-disabled")
                            if is_disabled == "true" or "disabled" in (await next_btn.get_attribute("class") or ""):
                                print("Đã đến trang cuối cùng.")
                                break
                            else:
                                print("Chuyển sang trang tiếp theo...")
                                await next_btn.click()
                                await page.wait_for_timeout(4000)
                                page_num += 1
                        else:
                            print("Không có phân trang.")
                            break
                            
                    await page.close()
            except Exception as e:
                print(f"Lỗi khi cào vbpl.vn: {e}")
                if job_id:
                    job = db.query(ScraperJob).filter(ScraperJob.id == job_id).first()
                    if job:
                        job.message = f"Lỗi vbpl: {str(e)}"
                        job.status = "FAILED"
                        db.commit()
            finally:
                await browser.close()
                db.close()

    async def scrape_moc(self, job_id: int = None, mode: str = "full"):
        """
        Cào dữ liệu từ Cổng thông tin Bộ Xây Dựng (moc.gov.vn)
        """
        db = SessionLocal()
        async with async_playwright() as p:
            browser = await p.chromium.launch(
                headless=True, 
                args=['--start-maximized', '--disable-blink-features=AutomationControlled'],
                proxy={
                    "server": "http://103.129.127.244:8088"
                }
            )
            context = await browser.new_context(
                no_viewport=True,
                user_agent="Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
                accept_downloads=True,
                ignore_https_errors=True
            )
            page = await context.new_page()
            try:
                print("Truy cập trang moc.gov.vn...")
                await page.goto("https://moc.gov.vn/vn/Pages/Vanbanphapluat.aspx", timeout=60000, wait_until="domcontentloaded")
                await page.wait_for_load_state("networkidle")
                
                page_num = 1
                while True:
                    if not await self._check_job_status(db, job_id):
                        break

                    if job_id:
                        job = db.query(ScraperJob).filter(ScraperJob.id == job_id).first()
                        if job:
                            job.current_page = page_num
                            db.commit()
                            
                    print(f"--- MOC: Đang cào trang {page_num} ---")
                    
                    # Lấy tất cả các thẻ tr trong bảng danh sách
                    rows = await page.locator(".table-bordered tbody tr").all()
                    if not rows:
                        rows = await page.locator("table tbody tr").all()
                        
                    print(f"Tìm thấy {len(rows)} hàng.")
                    
                    # Lặp qua các hàng, bỏ qua header (tuỳ web)
                    for row in rows:
                        if not await self._check_job_status(db, job_id):
                            break
                            
                        tds = await row.locator("td").all()
                        if len(tds) < 3:
                            continue # Thường là dòng header hoặc sai cấu trúc
                            
                        doc_number = (await tds[0].inner_text()).strip()
                        issue_date_raw = (await tds[1].inner_text()).strip()
                        title_el = tds[2].locator("a").first
                        if await title_el.count() == 0:
                            title_el = tds[2]
                        title = (await title_el.inner_text()).strip()
                        
                        if not doc_number or not title:
                            continue
                            
                        if "số/ký hiệu" in doc_number.lower() or "ngày ban hành" in issue_date_raw.lower():
                            continue
                            
                        doc_data = {
                            "title": title,
                            "origin_url": "https://moc.gov.vn",
                            "document_number": doc_number,
                            "status": "N/A",
                            "document_type": "N/A",
                            "signer": "N/A",
                            "issuing_body": "Bộ Xây Dựng",
                            "content": "",
                            "issue_date": "N/A",
                            "effective_date": "N/A"
                        }
                        
                        # Định dạng lại ngày
                        try:
                            dt = datetime.strptime(issue_date_raw, "%d/%m/%Y")
                            doc_data["issue_date"] = dt.strftime("%Y-%m-%d")
                        except ValueError:
                            doc_data["issue_date"] = issue_date_raw
                            
                        print(f"Đang mở chi tiết: {title}")
                        
                        try:
                            # Mở trang chi tiết
                            href = await title_el.get_attribute("href")
                            if href:
                                if href.startswith("http"):
                                    full_url = href
                                elif href.startswith("/"):
                                    full_url = f"https://moc.gov.vn{href}"
                                else:
                                    full_url = f"https://moc.gov.vn/vn/pages/{href}"
                                    
                                doc_data["origin_url"] = full_url
                                
                                new_page = await context.new_page()
                                await new_page.goto(full_url, timeout=60000, wait_until="domcontentloaded")
                                await new_page.wait_for_load_state("networkidle")
                                
                                # Lấy thêm thông tin từ bảng chi tiết trên trang MOC
                                try:
                                    detail_rows = await new_page.locator("table tr").all()
                                    for row in detail_rows:
                                        cells = await row.locator("td, th").all()
                                        if len(cells) >= 2:
                                            label = (await cells[0].inner_text()).strip().lower()
                                            val1 = (await cells[1].inner_text()).strip() if len(cells) > 1 else ""
                                            val2 = (await cells[2].inner_text()).strip() if len(cells) > 2 else ""
                                            
                                            if "ngày hiệu lực" in label:
                                                try:
                                                    dt = datetime.strptime(val1, "%d/%m/%Y")
                                                    doc_data["effective_date"] = dt.strftime("%Y-%m-%d")
                                                except:
                                                    doc_data["effective_date"] = val1
                                            elif "hình thức văn bản" in label:
                                                doc_data["document_type"] = val1
                                            elif "đơn vị ban hành" in label or "người ký" in label:
                                                if val1: doc_data["issuing_body"] = val1
                                                if val2: doc_data["signer"] = val2
                                except Exception as e:
                                    print(f"Lỗi phân tích bảng chi tiết MOC: {e}")
                                
                                # Tìm Tệp đính kèm
                                attachment_links = await new_page.locator("a[href*='.pdf'], a[href*='.doc'], a[href*='/TepDinhKem/']").all()
                                if not attachment_links:
                                    # Thử selector khác nếu MOC đổi link
                                    attachment_links = await new_page.locator("a:has-text('Tải về'), a:has-text('Download')").all()
                                
                                full_content = ""
                                for att_link in attachment_links:
                                    text_from_file = await self._download_and_extract_file(new_page, att_link)
                                    if text_from_file:
                                        full_content += f"\n\n--- Trích xuất từ tệp đính kèm ---\n\n{text_from_file}"
                                
                                if full_content:
                                    doc_data["content"] = full_content
                                else:
                                    # Thử lấy content text trực tiếp từ body nếu không có file
                                    try:
                                        body_content = await new_page.locator(".detail-content, .article-content").inner_text()
                                        doc_data["content"] = body_content
                                    except:
                                        doc_data["content"] = ""
                                        
                                await new_page.close()
                                
                        except Exception as ex:
                            print(f"Lỗi lấy chi tiết MOC: {ex}")
                            continue
                            
                        try:
                            result = self._save_document(db, doc_data, job_id)
                            if job_id:
                                job = db.query(ScraperJob).filter(ScraperJob.id == job_id).first()
                                if job:
                                    if result == "added":
                                        job.docs_added += 1
                                    elif result == "updated":
                                        job.docs_updated += 1
                                    db.commit()
                        except Exception as e:
                            print(f"Lỗi lưu DB (MOC): {e}")
                            db.rollback()

                    if not await self._check_job_status(db, job_id):
                        break

                    if mode in ["schedule_scrape", "manual_update", "schedule_update"] and page_num >= 2:
                        print("Chế độ LỊCH TRÌNH: Đã cào đủ số trang quy định.")
                        break

                    # Phân trang
                    # Tìm nút Next
                    next_btns = await page.locator("a.NextPage, a:has-text('Trang sau'), a[title='Trang sau']").all()
                    has_next = False
                    for nb in next_btns:
                        if await nb.is_visible():
                            await nb.click()
                            await page.wait_for_timeout(4000)
                            page_num += 1
                            has_next = True
                            break
                            
                    if not has_next:
                        print("Không tìm thấy trang tiếp theo hoặc đã hết.")
                        break

            except Exception as e:
                print(f"Lỗi moc: {str(e)}")
                if job_id:
                    job = db.query(ScraperJob).filter(ScraperJob.id == job_id).first()
                    if job:
                        job.message = f"Lỗi moc: {str(e)}"
                        job.status = "FAILED"
                        db.commit()
            finally:
                await browser.close()
                db.close()

    async def update_existing_documents(self, job_id: int = None):
        """
        Cập nhật lại trạng thái và nội dung cho các văn bản đang Có/Chưa hiệu lực
        """
        db = SessionLocal()
        # Lấy danh sách văn bản cần cập nhật (bỏ qua văn bản Hết hiệu lực, Ngưng hiệu lực)
        docs = db.query(Document).filter(
            Document.status.ilike("%còn hiệu lực%"),
            Document.status.ilike("%chưa có hiệu lực%")
        ).all()
        
        # Sửa lại điều kiện or_ cho đúng (cần import or_ từ sqlalchemy)
        # Tuy nhiên ta có thể filter bằng Python để đơn giản
        docs = db.query(Document).all()
        target_docs = [d for d in docs if d.status and d.status.lower() in ["còn hiệu lực", "chưa có hiệu lực", "n/a"]]
        
        print(f"Bắt đầu cập nhật {len(target_docs)} văn bản...")
        
        if job_id:
            job = db.query(ScraperJob).filter(ScraperJob.id == job_id).first()
            if job:
                job.total_pages = len(target_docs)
                db.commit()

        async with async_playwright() as p:
            browser = await p.chromium.launch(headless=True, args=['--disable-blink-features=AutomationControlled'])
            context = await browser.new_context(
                user_agent="Mozilla/5.0 (Windows NT 10.0; Win64; x64)",
                accept_downloads=True,
                ignore_https_errors=True
            )
            
            try:
                for idx, doc in enumerate(target_docs):
                    if not await self._check_job_status(db, job_id):
                        break
                        
                    if job_id:
                        job = db.query(ScraperJob).filter(ScraperJob.id == job_id).first()
                        if job:
                            job.current_page = idx + 1
                            db.commit()

                    url = doc.origin_url
                    if not url or "http" not in url:
                        continue
                        
                    print(f"Đang cập nhật [{idx+1}/{len(target_docs)}]: {doc.title}")
                    
                    doc_data = {
                        "title": doc.title,
                        "origin_url": url,
                        "document_number": doc.document_number,
                        "status": doc.status,
                        "document_type": doc.document_type,
                        "signer": doc.signer,
                        "issuing_body": doc.issuing_body,
                        "content": doc.content,
                        "issue_date": doc.issue_date.strftime("%Y-%m-%d") if doc.issue_date else "N/A",
                        "effective_date": doc.effective_date.strftime("%Y-%m-%d") if doc.effective_date else "N/A"
                    }
                    
                    try:
                        page = await context.new_page()
                        await page.goto(url, timeout=60000, wait_until="domcontentloaded")
                        await page.wait_for_load_state("networkidle")
                        await page.wait_for_timeout(2000)
                        
                        if "vbpl.vn" in url:
                            try:
                                await page.wait_for_selector(".ant-tabs-tabpane, #contentBody", timeout=15000)
                                content_el = page.locator(".ant-tabs-tabpane, #contentBody").first
                                doc_data["content"] = await content_el.inner_text()
                                
                                tab_btn = page.locator(".ant-tabs-tab-btn", has_text="Thuộc tính").first
                                if await tab_btn.count() > 0:
                                    await tab_btn.click()
                                    await page.wait_for_timeout(2000)
                                    props_el = page.locator(".ant-tabs-tabpane-active").first
                                    props_text = await props_el.inner_text()
                                    self._parse_properties(props_text, doc_data)
                            except Exception as ex:
                                print(f"Lỗi đọc nội dung vbpl: {ex}")
                                
                        elif "moc.gov.vn" in url:
                            try:
                                attachment_links = await page.locator("a[href*='/TepDinhKem/'], a:has-text('Tải về'), a:has-text('Tệp đính kèm')").all()
                                full_content = ""
                                for att_link in attachment_links:
                                    text_from_file = await self._download_and_extract_file(page, att_link)
                                    if text_from_file:
                                        full_content += f"\n\n--- Trích xuất từ tệp đính kèm ---\n\n{text_from_file}"
                                
                                if full_content:
                                    doc_data["content"] = full_content
                                else:
                                    try:
                                        body_content = await page.locator(".detail-content, .article-content").inner_text()
                                        doc_data["content"] = body_content
                                    except:
                                        pass
                            except Exception as ex:
                                print(f"Lỗi đọc nội dung moc: {ex}")
                                
                        await page.close()
                    except Exception as e:
                        print(f"Lỗi truy cập {url}: {e}")
                        continue
                        
                    # Lưu lại
                    try:
                        # Gọi _save_document. Nếu hash thay đổi (đặc biệt là metadata status) nó sẽ cập nhật
                        result = self._save_document(db, doc_data, job_id)
                        if job_id:
                            job = db.query(ScraperJob).filter(ScraperJob.id == job_id).first()
                            if job:
                                if result == "updated": job.docs_updated += 1
                                db.commit()
                    except Exception as e:
                        print(f"Lỗi lưu DB: {e}")

            except Exception as e:
                print(f"Lỗi update loop: {e}")
                if job_id:
                    job = db.query(ScraperJob).filter(ScraperJob.id == job_id).first()
                    if job:
                        job.message = f"Lỗi update loop: {str(e)}"
                        job.status = "FAILED"
                        db.commit()
            finally:
                await browser.close()
                db.close()

async def run_scraper_logic(job_id: int):
    """
    Hàm entrypoint cho background task.
    """
    db = SessionLocal()
    job = db.query(ScraperJob).filter(ScraperJob.id == job_id).first()
    if not job:
        db.close()
        return
        
    job.status = "RUNNING"
    db.commit()
    
    scraper = LuatXayDungScraper()
    
    try:
        # Nếu job là loại update hiện tại
        if job.mode in ["manual_update", "schedule_update"]:
            await scraper.update_existing_documents(job_id=job.id)
        else:
            # Chế độ cào thông thường (manual_scrape, schedule_scrape)
            if job.source in ["vbpl", "all"]:
                await scraper.scrape_vbpl(job_id=job.id, mode=job.mode)
                
            db.refresh(job)
            if job.status not in ["STOPPED", "FAILED"] and job.source in ["moc", "all"]:
                await scraper.scrape_moc(job_id=job.id, mode=job.mode)
            
        db.refresh(job)
        if job.status not in ["STOPPED", "FAILED"]:
            job.status = "COMPLETED"
    except Exception as e:
        error_msg = traceback.format_exc()
        print(f"Scraper Error: {error_msg}")
        job.status = "FAILED"
        job.message = str(e) if str(e) else "Lỗi không xác định (xem log terminal)"
    finally:
        job.ended_at = datetime.now()
        db.commit()
        db.close()

def start_scraper_background_task(job_id: int):
    """Được gọi từ API dưới dạng sync để chạy trong Thread riêng (giúp Playwright hoạt động trên Windows)"""
    asyncio.run(run_scraper_logic(job_id))
